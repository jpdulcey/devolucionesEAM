import io
import re
from datetime import datetime

import gspread
import pandas as pd
import streamlit as st
from docx import Document
from oauth2client.service_account import ServiceAccountCredentials

st.set_page_config(page_title="Calidad analistas", layout="wide")

# =========================
# CONFIGURACIÓN GENERAL
# =========================
PUNTAJE_BASE = 100
SHEET_ID = "1xrDybkfOPlH3fLHEedPQG77Sf3PfUQzC7wKXPTber_g"

# =========================
# HEADER
# =========================
col1, col2 = st.columns([5, 2])

with col1:
    st.markdown(
        """
        <h1 style='margin-bottom: 0;'>DEVOLUCIONES</h1>
        <hr style='margin-top: 5px;'>
        """,
        unsafe_allow_html=True
    )

with col2:
    st.image("logo_DANE.jpg", use_container_width=True)


# =========================
# FUNCIONES AUXILIARES
# =========================
def limpiar_nombre_archivo(texto):
    texto = str(texto).strip()
    texto = re.sub(r'[\\/*?:"<>|]', "", texto)
    texto = re.sub(r"\s+", "_", texto)
    return texto


def inicializar_estado_modulo(prefix):
    defaults = {
        f"{prefix}_evaluacion_finalizada": False,
        f"{prefix}_seleccionados_finales": [],
        f"{prefix}_puntaje_final_final": 100,
        f"{prefix}_decision_final": "",
        f"{prefix}_decision_usuario": "",
        f"{prefix}_registro_guardado": False,
        f"{prefix}_registro_error": None,
        f"{prefix}_registro_ya_guardado": False,
        f"{prefix}_accion_pendiente": "",
        f"{prefix}_dias_respuesta": 1,
    }

    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def resetear_estado_modulo(prefix):
    st.session_state[f"{prefix}_evaluacion_finalizada"] = False
    st.session_state[f"{prefix}_seleccionados_finales"] = []
    st.session_state[f"{prefix}_puntaje_final_final"] = 100
    st.session_state[f"{prefix}_decision_final"] = ""
    st.session_state[f"{prefix}_decision_usuario"] = ""
    st.session_state[f"{prefix}_registro_guardado"] = False
    st.session_state[f"{prefix}_registro_error"] = None
    st.session_state[f"{prefix}_registro_ya_guardado"] = False
    st.session_state[f"{prefix}_accion_pendiente"] = ""
    st.session_state[f"{prefix}_dias_respuesta"] = 1


def conectar_sheet():
    scope = [
        "https://spreadsheets.google.com/feeds",
        "https://www.googleapis.com/auth/drive",
    ]

    creds_dict = dict(st.secrets["gcp_service_account"])

    creds = ServiceAccountCredentials.from_json_keyfile_dict(
        creds_dict,
        scope
    )

    client = gspread.authorize(creds)
    spreadsheet = client.open_by_key(SHEET_ID)
    sheet = spreadsheet.sheet1
    return sheet


def guardar_registro_sheet(
    fecha,
    tipo,
    analista,
    territorial,
    nordemp,
    nordest,
    nomest,
    monitor,
    codsede,
    calificacion,
    resultado,
    dias
):
    try:
        sheet = conectar_sheet()

        sheet.append_row(
            [
                fecha,
                tipo,
                analista,
                territorial,
                nordemp,
                nordest,
                nomest,
                monitor,
                codsede,
                calificacion,
                resultado,
                dias
            ],
            value_input_option="USER_ENTERED"
        )

        return True, None

    except Exception as e:
        if str(e).strip() == "<Response [200]>":
            return True, None

        return False, repr(e)


@st.cache_data
def cargar_fuentes():
    df_control = pd.read_excel("control.xlsx")
    df_capdirest = pd.read_excel("capdirest.xlsx")

    df_control.columns = df_control.columns.str.strip().str.lower()
    df_capdirest.columns = df_capdirest.columns.str.strip().str.lower()

    df_control["nordest"] = df_control["nordest"].astype(str).str.strip()
    df_capdirest["nordest"] = df_capdirest["nordest"].astype(str).str.strip()

    if "usuario" not in df_control.columns:
        raise KeyError("No existe la columna 'usuario' en control.xlsx")
    if "usuarioss" not in df_control.columns:
        raise KeyError("No existe la columna 'usuarioss' en control.xlsx")
    if "codsede" not in df_control.columns:
        raise KeyError("No existe la columna 'codsede' en control.xlsx")
    if "nomest" not in df_capdirest.columns:
        raise KeyError("No existe la columna 'nomest' en capdirest.xlsx")
    if "nordemp" not in df_capdirest.columns:
        raise KeyError("No existe la columna 'nordemp' en capdirest.xlsx")

    df_control = df_control[["nordest", "usuario", "usuarioss", "codsede"]].drop_duplicates()
    df_control = df_control.rename(columns={
        "usuario": "analista",
        "usuarioss": "monitor",
        "codsede": "territorial"
    })

    df_capdirest = df_capdirest[["nordest", "nordemp", "nomest"]].drop_duplicates()

    df_base = df_control.merge(df_capdirest, on="nordest", how="outer")
    df_base = df_base.sort_values("nordest").reset_index(drop=True)

    return df_base


def preparar_df_puntajes(df):
    df.columns = [str(col).strip() for col in df.columns]

    columnas_requeridas = ["TÍTULO", "SUBTÍTULO_2", "PUNTAJE"]
    for col in columnas_requeridas:
        if col not in df.columns:
            raise KeyError(f"Falta la columna '{col}' en puntajes.xlsx")

    if "SUBTÍTULO_1" not in df.columns:
        df["SUBTÍTULO_1"] = ""

    df = df[["TÍTULO", "SUBTÍTULO_1", "SUBTÍTULO_2", "PUNTAJE"]].copy()
    df["orden"] = range(len(df))

    df["TÍTULO"] = df["TÍTULO"].fillna("").astype(str).str.strip()
    df["SUBTÍTULO_1"] = df["SUBTÍTULO_1"].fillna("").astype(str).str.strip()
    df["SUBTÍTULO_2"] = df["SUBTÍTULO_2"].fillna("").astype(str).str.strip()
    df["PUNTAJE"] = pd.to_numeric(df["PUNTAJE"], errors="coerce").fillna(0)

    df = df[(df["TÍTULO"] != "") & (df["SUBTÍTULO_2"] != "")].reset_index(drop=True)
    return df


@st.cache_data
def cargar_puntajes():
    hojas = pd.read_excel("puntajes.xlsx", sheet_name=None)

    if len(hojas) < 2:
        raise ValueError(
            "puntajes.xlsx debe tener al menos 2 hojas: una para Devoluciones y otra para Novedades."
        )

    nombres_hojas = list(hojas.keys())

    # Intenta tomar hojas por nombre; si no existen, toma las dos primeras
    hoja_dev = None
    hoja_nov = None

    for nombre in nombres_hojas:
        nombre_norm = str(nombre).strip().lower()
        if nombre_norm == "devoluciones":
            hoja_dev = nombre
        elif nombre_norm == "novedades":
            hoja_nov = nombre

    if hoja_dev is None:
        hoja_dev = nombres_hojas[0]

    if hoja_nov is None:
        hoja_nov = nombres_hojas[1]

    df_devoluciones = preparar_df_puntajes(hojas[hoja_dev])
    df_novedades = preparar_df_puntajes(hojas[hoja_nov])

    return df_devoluciones, df_novedades, hoja_dev, hoja_nov


def generar_word(
    tipo_proceso,
    nordemp,
    nordest,
    analista,
    monitor,
    territorial,
    establecimiento,
    seleccionados,
    puntaje_final,
    decision,
    dias=None
):
    doc = Document()

    doc.add_heading(f"Resultado de Evaluación - {tipo_proceso}", 1)

    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Tipo: {tipo_proceso}")
    doc.add_paragraph(f"NORDEMP: {nordemp}")
    doc.add_paragraph(f"NORDEST: {nordest}")
    doc.add_paragraph(f"Analista: {analista}")
    doc.add_paragraph(f"Monitor: {monitor}")
    doc.add_paragraph(f"Territorial: {territorial}")
    doc.add_paragraph(f"Establecimiento: {establecimiento}")

    doc.add_heading("Resumen", 2)
    doc.add_paragraph(f"Puntaje final: {puntaje_final}")
    doc.add_paragraph(f"Resultado: {decision}")

    if decision == "ENVIAR CORREO":
        doc.add_paragraph(f"Días otorgados para responder: {dias if dias is not None else ''}")

    doc.add_heading("Observaciones / Novedades registradas", 2)

    if not seleccionados:
        doc.add_paragraph("No se registraron observaciones.")
    else:
        seleccionados_ordenados = sorted(seleccionados, key=lambda x: x["orden"])
        modulo_actual = None
        categoria_actual = None

        for item in seleccionados_ordenados:
            if item["titulo"] != modulo_actual:
                doc.add_heading(item["titulo"], 3)
                modulo_actual = item["titulo"]
                categoria_actual = None

            if item["categoria"]:
                if item["categoria"] != categoria_actual:
                    doc.add_heading(item["categoria"], 4)
                    categoria_actual = item["categoria"]

            p_sub = doc.add_paragraph()
            run = p_sub.add_run(item["subtitulo"])
            run.bold = True

            if item["texto"]:
                doc.add_paragraph(item["texto"])
            else:
                doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def calcular_puntaje_final(seleccionados):
    total_descuento = sum(float(x["puntaje"]) for x in seleccionados)
    puntaje_final = PUNTAJE_BASE - total_descuento

    if puntaje_final < 0:
        puntaje_final = 0
    if puntaje_final > 100:
        puntaje_final = 100

    return puntaje_final


def render_modulo(
    df_puntajes,
    prefix,
    tipo_proceso,
    nordemp,
    nordest,
    analista,
    monitor,
    territorial,
    establecimiento,
    codsede
):
    inicializar_estado_modulo(prefix)

    st.subheader(f"2. {tipo_proceso.title()}")

    titulos_en_orden = df_puntajes["TÍTULO"].drop_duplicates().tolist()

    for titulo in titulos_en_orden:
        df_titulo = df_puntajes[df_puntajes["TÍTULO"] == titulo].copy()
        df_titulo = df_titulo.sort_values("orden")

        with st.expander(titulo, expanded=False):
            categorias = [
                x for x in df_titulo["SUBTÍTULO_1"].drop_duplicates().tolist()
                if str(x).strip() != ""
            ]

            if categorias:
                for categoria in categorias:
                    df_categoria = df_titulo[df_titulo["SUBTÍTULO_1"] == categoria].copy()

                    with st.expander(categoria, expanded=False):
                        for _, fila2 in df_categoria.iterrows():
                            idx = int(fila2["orden"])
                            check_key = f"{prefix}_check_{idx}"
                            text_key = f"{prefix}_text_{idx}"

                            st.checkbox(
                                f"{fila2['SUBTÍTULO_2']} ({fila2['PUNTAJE']})",
                                key=check_key
                            )

                            if st.session_state.get(check_key, False):
                                st.text_area(
                                    f"Redacción: {fila2['SUBTÍTULO_2']}",
                                    key=text_key,
                                    height=120,
                                    placeholder="Aquí el analista puede escribir o pegar la observación..."
                                )
            else:
                for _, fila2 in df_titulo.iterrows():
                    idx = int(fila2["orden"])
                    check_key = f"{prefix}_check_{idx}"
                    text_key = f"{prefix}_text_{idx}"

                    st.checkbox(
                        f"{fila2['SUBTÍTULO_2']} ({fila2['PUNTAJE']})",
                        key=check_key
                    )

                    if st.session_state.get(check_key, False):
                        st.text_area(
                            f"Redacción: {fila2['SUBTÍTULO_2']}",
                            key=text_key,
                            height=120,
                            placeholder="Aquí el analista puede escribir o pegar la observación..."
                        )

    st.divider()

    if st.button(f"Finalizar {tipo_proceso}", key=f"{prefix}_finalizar_btn"):
        seleccionados = []

        for _, fila2 in df_puntajes.sort_values("orden").iterrows():
            idx = int(fila2["orden"])
            check_key = f"{prefix}_check_{idx}"
            text_key = f"{prefix}_text_{idx}"

            if st.session_state.get(check_key, False):
                seleccionados.append({
                    "titulo": fila2["TÍTULO"],
                    "categoria": fila2["SUBTÍTULO_1"],
                    "subtitulo": fila2["SUBTÍTULO_2"],
                    "puntaje": fila2["PUNTAJE"],
                    "texto": st.session_state.get(text_key, "").strip(),
                    "orden": idx
                })

        puntaje_final = calcular_puntaje_final(seleccionados)
        recomendacion = "DEVOLVER FUENTE" if puntaje_final < 90 else "ENVIAR CORREO"

        st.session_state[f"{prefix}_seleccionados_finales"] = seleccionados
        st.session_state[f"{prefix}_puntaje_final_final"] = puntaje_final
        st.session_state[f"{prefix}_decision_final"] = recomendacion
        st.session_state[f"{prefix}_decision_usuario"] = ""
        st.session_state[f"{prefix}_evaluacion_finalizada"] = True
        st.session_state[f"{prefix}_registro_guardado"] = False
        st.session_state[f"{prefix}_registro_error"] = None
        st.session_state[f"{prefix}_registro_ya_guardado"] = False
        st.session_state[f"{prefix}_accion_pendiente"] = ""
        st.session_state[f"{prefix}_dias_respuesta"] = 1

    if st.session_state[f"{prefix}_evaluacion_finalizada"]:
        st.subheader("3. Resultado")

        puntaje = st.session_state[f"{prefix}_puntaje_final_final"]
        recomendacion = st.session_state[f"{prefix}_decision_final"]
        decision_confirmada = st.session_state[f"{prefix}_registro_ya_guardado"]

        c1, c2 = st.columns(2)
        c1.metric("Puntaje final", f"{puntaje:g}")
        c2.metric("Recomendación del sistema", recomendacion)

        if recomendacion == "DEVOLVER FUENTE":
            st.error("Según la validación, se recomienda devolver la fuente. Seleccione la acción a realizar:")
        else:
            st.success("Según la validación, se recomienda enviar correo. Seleccione la acción a realizar:")

        col_btn1, col_btn2 = st.columns(2)

        with col_btn1:
            if recomendacion == "ENVIAR CORREO":
                if st.button(
                    "✅ Enviar correo",
                    key=f"{prefix}_btn_correo",
                    type="primary",
                    disabled=decision_confirmada
                ):
                    st.session_state[f"{prefix}_accion_pendiente"] = "ENVIAR CORREO"
            else:
                if st.button(
                    "✅ Enviar correo",
                    key=f"{prefix}_btn_correo_alt",
                    disabled=decision_confirmada
                ):
                    st.session_state[f"{prefix}_accion_pendiente"] = "ENVIAR CORREO"

        with col_btn2:
            if recomendacion == "DEVOLVER FUENTE":
                if st.button(
                    "🔁 Devolver fuente",
                    key=f"{prefix}_btn_devolver",
                    type="primary",
                    disabled=decision_confirmada
                ):
                    st.session_state[f"{prefix}_accion_pendiente"] = "DEVOLVER FUENTE"
            else:
                if st.button(
                    "🔁 Devolver fuente",
                    key=f"{prefix}_btn_devolver_alt",
                    disabled=decision_confirmada
                ):
                    st.session_state[f"{prefix}_accion_pendiente"] = "DEVOLVER FUENTE"

        if st.session_state[f"{prefix}_accion_pendiente"] and not decision_confirmada:
            st.divider()
            st.subheader("Confirmación")

            if st.session_state[f"{prefix}_accion_pendiente"] == "ENVIAR CORREO":
                st.info("¿Confirma que enviará el correo?")

                dias = st.number_input(
                    "Días otorgados para responder el correo",
                    min_value=1,
                    step=1,
                    format="%d",
                    key=f"{prefix}_dias_input_correo"
                )

                c1, c2 = st.columns(2)

                with c1:
                    if st.button("Confirmar envío", key=f"{prefix}_confirmar_envio", type="primary"):
                        st.session_state[f"{prefix}_decision_usuario"] = "ENVIAR CORREO"
                        st.session_state[f"{prefix}_dias_respuesta"] = int(dias)

                        fecha_registro = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                        ok, error_msg = guardar_registro_sheet(
                            fecha=fecha_registro,
                            tipo=tipo_proceso.upper(),
                            analista=str(analista),
                            territorial=str(territorial),
                            nordemp=str(nordemp),
                            nordest=str(nordest),
                            nomest=str(establecimiento),
                            monitor=str(monitor),
                            codsede=str(codsede),
                            calificacion=float(puntaje),
                            resultado="ENVIAR CORREO",
                            dias=int(dias)
                        )

                        st.session_state[f"{prefix}_registro_guardado"] = ok
                        st.session_state[f"{prefix}_registro_error"] = error_msg
                        st.session_state[f"{prefix}_registro_ya_guardado"] = True
                        st.session_state[f"{prefix}_accion_pendiente"] = ""

                with c2:
                    if st.button("Volver", key=f"{prefix}_volver_correo"):
                        st.session_state[f"{prefix}_accion_pendiente"] = ""

            elif st.session_state[f"{prefix}_accion_pendiente"] == "DEVOLVER FUENTE":
                st.warning("¿Confirma que devolverá la fuente?")

                c1, c2 = st.columns(2)

                with c1:
                    if st.button("Confirmar devolución", key=f"{prefix}_confirmar_devolucion", type="primary"):
                        st.session_state[f"{prefix}_decision_usuario"] = "DEVOLVER FUENTE"
                        st.session_state[f"{prefix}_dias_respuesta"] = ""

                        fecha_registro = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                        ok, error_msg = guardar_registro_sheet(
                            fecha=fecha_registro,
                            tipo=tipo_proceso.upper(),
                            analista=str(analista),
                            territorial=str(territorial),
                            nordemp=str(nordemp),
                            nordest=str(nordest),
                            nomest=str(establecimiento),
                            monitor=str(monitor),
                            codsede=str(codsede),
                            calificacion=float(puntaje),
                            resultado="DEVOLVER FUENTE",
                            dias=""
                        )

                        st.session_state[f"{prefix}_registro_guardado"] = ok
                        st.session_state[f"{prefix}_registro_error"] = error_msg
                        st.session_state[f"{prefix}_registro_ya_guardado"] = True
                        st.session_state[f"{prefix}_accion_pendiente"] = ""

                with c2:
                    if st.button("Volver", key=f"{prefix}_volver_devolucion"):
                        st.session_state[f"{prefix}_accion_pendiente"] = ""

        if st.session_state.get(f"{prefix}_decision_usuario"):
            st.divider()
            st.subheader("Decisión final seleccionada")

            if st.session_state[f"{prefix}_decision_usuario"] == "DEVOLVER FUENTE":
                st.error("DEVOLVER FUENTE")
            else:
                st.success(
                    f"ENVIAR CORREO | Días otorgados: {st.session_state[f'{prefix}_dias_respuesta']}"
                )

            if st.session_state[f"{prefix}_registro_guardado"]:
                st.success("Registro guardado en Google Sheets con la decisión final del usuario.")
            elif st.session_state[f"{prefix}_registro_error"]:
                st.error(f"No se pudo guardar en Google Sheets: {st.session_state[f'{prefix}_registro_error']}")

            archivo = generar_word(
                tipo_proceso=tipo_proceso.upper(),
                nordemp=nordemp,
                nordest=nordest,
                analista=analista,
                monitor=monitor,
                territorial=territorial,
                establecimiento=establecimiento,
                seleccionados=st.session_state[f"{prefix}_seleccionados_finales"],
                puntaje_final=puntaje,
                decision=st.session_state[f"{prefix}_decision_usuario"],
                dias=st.session_state[f"{prefix}_dias_respuesta"] if st.session_state[f"{prefix}_decision_usuario"] == "ENVIAR CORREO" else None
            )

            nombre_archivo = (
                f"{limpiar_nombre_archivo(tipo_proceso)}_"
                f"{limpiar_nombre_archivo(nordemp)}_"
                f"{limpiar_nombre_archivo(nordest)}_"
                f"{limpiar_nombre_archivo(establecimiento)}.docx"
            )

            st.download_button(
                "Descargar Word",
                archivo,
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"{prefix}_download_word"
            )


# =========================
# CARGA DE DATOS
# =========================
df_base = cargar_fuentes()
df_devoluciones, df_novedades, hoja_dev, hoja_nov = cargar_puntajes()

# Inicializar estados de ambos módulos
inicializar_estado_modulo("dev")
inicializar_estado_modulo("nov")

# =========================
# IDENTIFICACIÓN
# =========================
st.subheader("1. Identificación")

modo = st.radio("Modo de búsqueda", ["Escribir", "Seleccionar"], horizontal=True)

nordest = ""
if modo == "Escribir":
    nordest = st.text_input("NORDEST").strip()
else:
    lista = [""] + df_base["nordest"].dropna().astype(str).tolist()
    nordest = st.selectbox("NORDEST", lista)

if nordest:
    fila = df_base[df_base["nordest"] == nordest]

    if not fila.empty:
        analista = fila.iloc[0]["analista"]
        monitor = fila.iloc[0]["monitor"]
        territorial = fila.iloc[0]["territorial"]
        establecimiento = fila.iloc[0]["nomest"]
        nordemp = fila.iloc[0]["nordemp"]

        codsede = territorial

        col1, col2, col3, col4, col5 = st.columns(5)

        with col1:
            st.text_input("NORDEMP", value=str(nordemp), disabled=True)
        with col2:
            st.text_input("Analista", value=str(analista), disabled=True)
        with col3:
            st.text_input("Monitor", value=str(monitor), disabled=True)
        with col4:
            st.text_input("Territorial", value=str(territorial), disabled=True)
        with col5:
            st.text_input("Establecimiento", value=str(establecimiento), disabled=True)

        st.divider()

        tab1, tab2 = st.tabs(["Devoluciones", "Novedades"])

        with tab1:
            render_modulo(
                df_puntajes=df_devoluciones,
                prefix="dev",
                tipo_proceso="Devoluciones",
                nordemp=nordemp,
                nordest=nordest,
                analista=analista,
                monitor=monitor,
                territorial=territorial,
                establecimiento=establecimiento,
                codsede=codsede
            )

        with tab2:
            render_modulo(
                df_puntajes=df_novedades,
                prefix="nov",
                tipo_proceso="Novedades",
                nordemp=nordemp,
                nordest=nordest,
                analista=analista,
                monitor=monitor,
                territorial=territorial,
                establecimiento=establecimiento,
                codsede=codsede
            )

    else:
        st.warning("No se encontró ese NORDEST.")
